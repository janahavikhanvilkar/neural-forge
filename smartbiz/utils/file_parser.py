import os
from pathlib import Path
import re

def extract_text_from_file(file_path: str) -> str:
    """Extracts raw text content from PDF, DOCX, TXT, or Image files."""
    path = Path(file_path)
    if not path.exists():
        return ""
    
    ext = path.suffix.lower().replace('.', '')
    
    if ext == 'txt':
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            return f"Error reading TXT: {str(e)}"
            
    elif ext == 'pdf':
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            text_pages = []
            for i, page in enumerate(reader.pages):
                extracted = page.extract_text()
                if extracted:
                    text_pages.append(extracted)
            return "\n\n".join(text_pages)
        except Exception as e:
            return f"Error extracting PDF text: {str(e)}"
            
    elif ext in ('docx', 'doc'):
        try:
            import docx
            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n".join(paragraphs)
        except Exception as e:
            return f"Error extracting DOCX text: {str(e)}"
            
    elif ext in ('png', 'jpg', 'jpeg'):
        try:
            from PIL import Image
            img = Image.open(str(path))
            # If tesseract is available, use pytesseract, otherwise generate simulated OCR content based on image metadata or file name
            try:
                import pytesseract
                return pytesseract.image_to_string(img)
            except Exception:
                # Basic OCR fallback when tesseract binary is not installed
                return f"[Image File Processed: {path.name}, Size: {img.size[0]}x{img.size[1]}px]"
        except Exception as e:
            return f"Error processing image: {str(e)}"
            
    return ""

def clean_extracted_text(text: str) -> str:
    """Cleans up raw extracted text for AI consumption."""
    if not text:
        return ""
    # Normalize multiple whitespace/newlines
    cleaned = re.sub(r'[ \t]+', ' ', text)
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    return cleaned.strip()
